import logging
import tempfile
from datetime import datetime, date as date_type

import requests
from docx import Document
from fastapi import APIRouter, Depends, Request, HTTPException, Query, Body
from fastapi.responses import HTMLResponse, FileResponse
from sqlalchemy import func
from sqlalchemy.orm import Session
from openpyxl import Workbook

from app.core.config import settings
from app.db.session import get_db
from app.models.report import Report
from app.core.templates import templates
from app.core.security import admin_required
from app.models.user import User
from pydantic import BaseModel
from app.core.security import normalize_login

router = APIRouter()

logger = logging.getLogger(__name__)


@router.get("/admin", response_class=HTMLResponse)
async def admin_page(
    request: Request,
    date_from: str | None = Query(None, alias="from"),
    date_to: str | None = Query(None, alias="to"),
    db: Session = Depends(get_db),
    _=Depends(admin_required),
):
    """Render admin dashboard with reports list."""
    query = db.query(Report)
    if date_from:
        try:
            start_dt = datetime.strptime(date_from, "%Y-%m-%d")
            query = query.filter(Report.created_at >= start_dt)
        except ValueError:
            pass

    if date_to:
        try:
            end_dt = datetime.strptime(date_to, "%Y-%m-%d").replace(hour=23, minute=59, second=59)
            query = query.filter(Report.created_at <= end_dt)
        except ValueError:
            pass

    # Сортируем по дате создания (самые новые сверху)
    query = query.order_by(Report.created_at.desc())

    reports = query.all()
    return templates.TemplateResponse(
        "admin.html",
        {
            "request": request,
            "reports": reports,
            "from": date_from,
            "to": date_to,
        },
    )


@router.get("/tracker/queues")
async def list_queues(_=Depends(admin_required)):
    """Вернуть список очередей, доступных пользователю в Яндекс.Трекере."""
    headers = {
        "Authorization": f"OAuth {settings.TRACKER_TOKEN}",
        "X-Org-ID": settings.TRACKER_ORG_ID,
        "Content-Type": "application/json",
    }

    resp = requests.get("https://api.tracker.yandex.net/v3/queues", headers=headers)
    if resp.status_code != 200:
        raise HTTPException(status_code=resp.status_code, detail=resp.text)

    # Оставляем только основные поля
    data = [{"key": q["key"], "name": q.get("name") or q.get("display")} for q in resp.json()]
    return data


class CreateTasksRequest(BaseModel):
    queue: str
    assignees: list[str]


@router.post("/admin/create-tasks")
async def create_tasks_in_tracker(
    payload: CreateTasksRequest = Body(...),
    db: Session = Depends(get_db),
    _=Depends(admin_required),
):
    """Create predefined tasks in Yandex Tracker for given assignees."""

    queue = payload.queue or settings.TRACKER_QUEUE
    assignees = payload.assignees or ["yakovleva.sv"]
    task_texts = [
        "Создание отчета по поддержанным программам",
        "Анализ новых образовательных проектов",
        "Оценка привлечения новых ученых",
        "Подготовка презентации по итогам",
        "Формирование базы данных отчетов",
        "Проверка данных по поддержанным программам",
        "Актуализация информации об образовательных проектах",
        "Разработка рекомендаций для улучшения программ",
        "Организация встречи по отчетам",
        "Обновление внутреннего портала отчетности",
        "Отправка итогового отчета в отдел",
    ]

    headers = {
        "Authorization": f"OAuth {settings.TRACKER_TOKEN}",
        "X-Org-ID": settings.TRACKER_ORG_ID,
        "Content-Type": "application/json",
    }

    created_issues: list[str] = []
    for i, text in enumerate(task_texts):
        assignee = assignees[i % len(assignees)]
        issue_payload = {
            "queue": queue,
            "summary": text,
            "description": f"Задача: {text}\n\nПройдите опрос: http://localhost:8000/dashboard/{assignee}",
            "type": "task",
            "priority": {"id": "3"},
            "assignee": assignee,
        }
        resp = requests.post("https://api.tracker.yandex.net/v3/issues/", headers=headers, json=issue_payload)
        if resp.status_code == 201:
            created_issue = resp.json()
            issue_key = created_issue["key"]
            created_issues.append(issue_key)

            # Move issue to "В работу" status if transition available
            transitions_url = f"https://api.tracker.yandex.net/v3/issues/{issue_key}/transitions"
            trans_resp = requests.get(transitions_url, headers=headers)
            if trans_resp.status_code == 200:
                transitions = trans_resp.json()
                transition = next((t for t in transitions if t["display"].lower() == "в работу"), None)
                if transition:
                    exec_url = (
                        f"https://api.tracker.yandex.net/v3/issues/{issue_key}/transitions/{transition['id']}/_execute"
                    )
                    requests.post(exec_url, headers=headers, json={"comment": "Статус установлен автоматически"})
        else:
            logger.error("Ошибка создания задачи: %s", resp.text)

    return {
        "queue": queue,
        "created": len(created_issues),
        "issues": created_issues,
    }


@router.get("/admin/export/word")
def export_word(
    date_from: date_type | None = Query(None, alias="from"),
    date_to: date_type | None = Query(None, alias="to"),
    db: Session = Depends(get_db),
    _=Depends(admin_required),
):
    """Export reports to a Word (.docx) file."""
    query = db.query(Report)
    if date_from:
        query = query.filter(func.date(Report.created_at) >= date_from)
    if date_to:
        query = query.filter(func.date(Report.created_at) <= date_to)
    reports = query.all()

    doc = Document()
    doc.add_heading("Отчеты пользователей", 0)

    for report in reports:
        doc.add_paragraph(
            f"ID: {report.id}\nПользователь: {report.username}\nПрограмм поддержано: {report.programs_supported}\n"
            f"Проектов в программе: {report.projects_in_program}\nНовых ученых: {report.new_scientists_employed}\n",
            style="Normal",
        )
        doc.add_paragraph("------------------------------")

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        if date_from or date_to:
            fname_range = f"{date_from or 'all'}_{date_to or 'all'}"
        else:
            fname_range = "all"
        filename = f"reports_{fname_range}.docx"
        return FileResponse(
            path=tmp.name,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


@router.get("/admin/export/excel")
def export_excel(
    date_from: date_type | None = Query(None, alias="from"),
    date_to: date_type | None = Query(None, alias="to"),
    db: Session = Depends(get_db),
    _=Depends(admin_required),
):
    """Export reports to an Excel (.xlsx) file."""

    query = db.query(Report)
    if date_from:
        query = query.filter(func.date(Report.created_at) >= date_from)
    if date_to:
        query = query.filter(func.date(Report.created_at) <= date_to)
    reports = query.all()

    wb = Workbook()
    ws = wb.active  # type: ignore[assignment]
    ws.title = "Отчёты"  # type: ignore[attr-defined]

    # header
    headers = [
        "ID",
        "Пользователь",
        "Поддержано программ",
        "Проектов в программе",
        "Новых учёных",
        "Дата создания",
    ]
    ws.append(headers)  # type: ignore[arg-type]

    for rpt in reports:
        ws.append([  # type: ignore[arg-type]
            rpt.id,
            rpt.username,
            rpt.programs_supported,
            rpt.projects_in_program,
            rpt.new_scientists_employed,
            rpt.created_at.strftime("%Y-%m-%d %H:%M"),
        ])

    # autosize columns
    for col in ws.columns:  # type: ignore[attr-defined]
        max_len = max(len(str(cell.value)) if cell.value is not None else 0 for cell in col)
        ws.column_dimensions[col[0].column_letter].width = max_len + 2  # type: ignore[attr-defined]

    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
        wb.save(tmp.name)
        if date_from or date_to:
            fname_range = f"{date_from or 'all'}_{date_to or 'all'}"
        else:
            fname_range = "all"
        filename = f"reports_{fname_range}.xlsx"
        return FileResponse(
            path=tmp.name,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )


@router.get("/admin/users")
async def list_users(db: Session = Depends(get_db), _ = Depends(admin_required)):
    return [u.login for u in db.query(User).all()] 


@router.get("/tracker/queues/{queue_key}/users")
async def list_queue_users(queue_key: str, _=Depends(admin_required)):
    """Вернуть список логинов пользователей (teamUsers) указанной очереди."""

    headers = {
        "Authorization": f"OAuth {settings.TRACKER_TOKEN}",
        "X-Org-ID": settings.TRACKER_ORG_ID,
    }

    # Берём параметры очереди с teamUsers через expand.
    url = f"https://api.tracker.yandex.net/v3/queues/{queue_key}?expand=teamUsers"
    resp = requests.get(url, headers=headers)
    if resp.status_code != 200:
        # иногда ?expand может не поддерживаться; пробуем expand=all
        if resp.status_code == 400 or resp.status_code == 404:
            resp2 = requests.get(
                f"https://api.tracker.yandex.net/v3/queues/{queue_key}?expand=all",
                headers=headers,
            )
            resp = resp2
        if resp.status_code != 200:
            raise HTTPException(status_code=resp.status_code, detail=resp.text)

    data = resp.json()
    team = data.get("teamUsers", [])

    items: list[dict] = []
    for user in team:
        # Извлекаем строковый логин; сначала явное поле `login`, затем `uid`, и в крайнем случае `id` (цифры)
        raw_login = user.get("login") or user.get("uid") or user.get("id") or user.get("display")
        if not raw_login:
            continue
        login: str | None = None
        display_name = user.get("display") or raw_login

        # Если raw_login — числовой идентификатор без имени, пробуем запросить данные пользователя
        if isinstance(raw_login, int) or (isinstance(raw_login, str) and raw_login.isdigit()):
            user_detail = requests.get(
                f"https://api.tracker.yandex.net/v3/users/{raw_login}", headers=headers
            )
            if user_detail.status_code == 200:
                raw_login = user_detail.json().get("login") or user_detail.json().get("uid") or raw_login

        login = normalize_login(str(raw_login))
        items.append({"login": login, "display": display_name})

    items.sort(key=lambda x: x["display"].lower())
    return items 